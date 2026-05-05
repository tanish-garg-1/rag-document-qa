"""
TEST 16: UPLOAD ENDPOINT
Tests POST /upload with TXT, PDF, DOCX, and unsupported files.
REQUIRES GEMINI_API_KEY (for embeddings).
Run: python tests/test_16_upload_endpoint.py
"""

import sys
import os
import io
import tempfile
import shutil

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from dotenv import load_dotenv
load_dotenv()

print("=" * 60)
print("TEST 16: UPLOAD ENDPOINT")
print("=" * 60)

# --- 1. Check API key ---
print("\n--- API Key Check ---")
gemini_key = os.getenv("GEMINI_API_KEY")
if not gemini_key:
    print("  [SKIP] GEMINI_API_KEY not set")
    print("\n" + "=" * 60)
    print("TEST 16 SKIPPED (no GEMINI_API_KEY)")
    print("=" * 60)
    sys.exit(0)
print(f"  [PASS] GEMINI_API_KEY set")

# --- 2. Import and create client ---
print("\n--- Setup TestClient ---")
try:
    from fastapi.testclient import TestClient
    from app.main import app
    client = TestClient(app, raise_server_exceptions=False)
    print("  [PASS] TestClient ready")
except Exception as e:
    print(f"  [FAIL] Setup error: {e}")
    import traceback
    traceback.print_exc()
    sys.exit(1)

from app.utils.constants import FAISS_INDEX_DIR
root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
faiss_path = os.path.join(root, FAISS_INDEX_DIR)

# Backup index
backup_path = faiss_path + "_backup_test16"
if os.path.exists(faiss_path):
    if os.path.exists(backup_path):
        shutil.rmtree(backup_path)
    shutil.copytree(faiss_path, backup_path)
    shutil.rmtree(faiss_path)
    os.makedirs(faiss_path)

# --- 3. Upload TXT file ---
print("\n--- Upload TXT File ---")
txt_content = b"""Introduction to Artificial Intelligence

Artificial intelligence (AI) refers to the simulation of human intelligence in machines that are programmed to think like humans and mimic their actions. The term may also be applied to any machine that exhibits traits associated with a human mind such as learning and problem-solving.

The ideal characteristic of artificial intelligence is its ability to rationalize and take actions that have the best chance of achieving a specific goal. A subset of artificial intelligence is machine learning, which refers to the concept that computer programs can automatically learn from and adapt to new data without being assisted by humans.

Deep learning techniques enable this automatic learning through the absorption of huge amounts of unstructured data such as text, images, or video."""

try:
    import time
    t0 = time.time()
    response = client.post(
        "/upload",
        files=[("files", ("test_document.txt", io.BytesIO(txt_content), "text/plain"))],
    )
    elapsed = time.time() - t0
    print(f"  Upload time: {elapsed:.2f}s")
    print(f"  Status code: {response.status_code}")

    if response.status_code == 200:
        data = response.json()
        print(f"  Response: {data}")
        results = data.get("results", [])
        if results:
            r = results[0]
            print(f"  File: {r.get('filename')}")
            print(f"  Status: {r.get('status')}")
            print(f"  Chunks added: {r.get('chunks_added', 'N/A')}")
            if r.get("status") == "success":
                print("  [PASS] TXT file uploaded successfully")
            else:
                print(f"  [FAIL] Upload failed: {r.get('reason', 'unknown')}")
        else:
            print(f"  [FAIL] No results in response: {data}")
    else:
        print(f"  [FAIL] Upload status {response.status_code}: {response.text[:300]}")
        import traceback

except Exception as e:
    print(f"  [FAIL] Upload TXT error: {e}")
    import traceback
    traceback.print_exc()

# --- 4. Upload PDF file ---
print("\n--- Upload PDF File ---")
try:
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 72), "This is a test PDF for upload endpoint testing. Machine learning is fascinating.", fontsize=12)
    pdf_bytes = doc.tobytes()
    doc.close()

    t0 = time.time()
    response = client.post(
        "/upload",
        files=[("files", ("test_document.pdf", io.BytesIO(pdf_bytes), "application/pdf"))],
    )
    elapsed = time.time() - t0
    print(f"  Upload time: {elapsed:.2f}s")
    print(f"  Status code: {response.status_code}")

    if response.status_code == 200:
        data = response.json()
        results = data.get("results", [])
        if results:
            r = results[0]
            print(f"  Status: {r.get('status')}, chunks: {r.get('chunks_added', 'N/A')}")
            if r.get("status") == "success":
                print("  [PASS] PDF uploaded successfully")
            else:
                print(f"  [FAIL] PDF upload failed: {r.get('reason')}")
    else:
        print(f"  [FAIL] PDF upload status {response.status_code}: {response.text[:300]}")

except ImportError:
    print("  [SKIP] PyMuPDF not available for PDF test")
except Exception as e:
    print(f"  [FAIL] Upload PDF error: {e}")
    import traceback
    traceback.print_exc()

# --- 5. Upload DOCX file ---
print("\n--- Upload DOCX File ---")
try:
    from docx import Document
    doc = Document()
    doc.add_paragraph("Testing DOCX upload to RAG system.")
    doc.add_paragraph("Vector embeddings enable semantic search.")
    doc.add_paragraph("FAISS provides efficient similarity search.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_bytes = buf.read()

    t0 = time.time()
    response = client.post(
        "/upload",
        files=[("files", ("test_document.docx", io.BytesIO(docx_bytes),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
    )
    elapsed = time.time() - t0
    print(f"  Upload time: {elapsed:.2f}s")
    print(f"  Status code: {response.status_code}")

    if response.status_code == 200:
        data = response.json()
        results = data.get("results", [])
        if results:
            r = results[0]
            print(f"  Status: {r.get('status')}, chunks: {r.get('chunks_added', 'N/A')}")
            if r.get("status") == "success":
                print("  [PASS] DOCX uploaded successfully")
            else:
                print(f"  [FAIL] DOCX upload failed: {r.get('reason')}")
    else:
        print(f"  [FAIL] DOCX upload status {response.status_code}: {response.text[:300]}")

except ImportError:
    print("  [SKIP] python-docx not available")
except Exception as e:
    print(f"  [FAIL] Upload DOCX error: {e}")

# --- 6. Upload unsupported file ---
print("\n--- Upload Unsupported File Type ---")
try:
    response = client.post(
        "/upload",
        files=[("files", ("data.csv", io.BytesIO(b"col1,col2\n1,2\n3,4"), "text/csv"))],
    )
    print(f"  Status code: {response.status_code}")
    if response.status_code == 200:
        data = response.json()
        results = data.get("results", [])
        if results:
            r = results[0]
            print(f"  Status: {r.get('status')}, reason: {r.get('reason', 'N/A')}")
            if r.get("status") != "success":
                print("  [PASS] Unsupported file type correctly rejected")
            else:
                print("  [WARN] Unsupported file accepted as success")
    else:
        print(f"  Response: {response.text[:200]}")
except Exception as e:
    print(f"  [FAIL] Unsupported file test error: {e}")

# --- 7. Upload multiple files at once ---
print("\n--- Upload Multiple Files ---")
try:
    files = [
        ("files", ("multi1.txt", io.BytesIO(b"First file content about neural networks."), "text/plain")),
        ("files", ("multi2.txt", io.BytesIO(b"Second file content about deep learning."), "text/plain")),
    ]
    t0 = time.time()
    response = client.post("/upload", files=files)
    elapsed = time.time() - t0
    print(f"  Upload time: {elapsed:.2f}s")
    print(f"  Status code: {response.status_code}")

    if response.status_code == 200:
        data = response.json()
        results = data.get("results", [])
        print(f"  Results count: {len(results)}")
        if len(results) == 2:
            print("  [PASS] Both files processed")
        else:
            print(f"  [WARN] Expected 2 results, got {len(results)}")
        for r in results:
            status = r.get("status")
            fname = r.get("filename")
            chunks = r.get("chunks_added", "N/A")
            print(f"    {fname}: {status}, {chunks} chunks")
    else:
        print(f"  [FAIL] Multi-file upload status {response.status_code}")
except Exception as e:
    print(f"  [FAIL] Multi-file upload error: {e}")

# --- 8. POST /clear endpoint ---
print("\n--- POST /clear ---")
try:
    response = client.post("/clear")
    print(f"  Status code: {response.status_code}")
    if response.status_code == 200:
        print("  [PASS] /clear returned 200")
        print(f"  Response: {response.json()}")
    else:
        print(f"  [WARN] /clear status {response.status_code}: {response.text[:100]}")
except Exception as e:
    print(f"  [FAIL] /clear error: {e}")

# --- Restore index ---
print("\n--- Restore ---")
try:
    if os.path.exists(faiss_path):
        shutil.rmtree(faiss_path)
    if os.path.exists(backup_path):
        shutil.copytree(backup_path, faiss_path)
        shutil.rmtree(backup_path)
        print("  [INFO] Original index restored")
    else:
        os.makedirs(faiss_path, exist_ok=True)
        print("  [INFO] Fresh index directory created")
except Exception as e:
    print(f"  [WARN] Restore error: {e}")

print("\n" + "=" * 60)
print("TEST 16 COMPLETE")
print("=" * 60)
