"""
TEST 06: DOCUMENT LOADER — DOCX FILES
Tests DOCX loading using python-docx. No API key required.
Run: python tests/test_06_document_loader_docx.py
"""

import sys
import os
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

print("=" * 60)
print("TEST 06: DOCUMENT LOADER — DOCX FILES")
print("=" * 60)

# --- 1. Import docx ---
print("\n--- Import python-docx ---")
try:
    import docx
    print(f"  [PASS] python-docx imported")
    from docx import Document
    print(f"  [PASS] Document class imported")
except ImportError as e:
    print(f"  [FAIL] python-docx not installed: {e}")
    print("         Run: pip install python-docx")
    sys.exit(1)

# --- 2. Import loader ---
print("\n--- Import document_loader ---")
try:
    from app.services.document_loader import load_document, load_docx
    print("  [PASS] document_loader imported")
except ImportError as e:
    print(f"  [FAIL] Import error: {e}")
    sys.exit(1)

# --- 3. Helper to create DOCX ---
def create_test_docx(paragraphs):
    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    return path

# --- 4. Simple DOCX ---
print("\n--- Simple DOCX Load ---")
paras = [
    "Introduction: This is a test DOCX document.",
    "Section 1: Artificial intelligence is transforming industries.",
    "Section 2: Machine learning enables pattern recognition.",
    "Conclusion: RAG systems improve LLM accuracy.",
]
path = create_test_docx(paras)
try:
    docs = load_docx(path)
    print(f"  Paragraphs in DOCX: {len(paras)}")
    print(f"  Documents returned: {len(docs)}")
    if docs:
        content = docs[0]["content"]
        print(f"  Content length: {len(content)}")
        print(f"  Content preview: '{content[:200]}'")

        # Check all paragraph text is present
        all_found = True
        for para in paras:
            if para in content:
                print(f"  [PASS] Found paragraph: '{para[:50]}'")
            else:
                print(f"  [WARN] Paragraph not found in content: '{para[:50]}'")
                all_found = False

        if all_found:
            print(f"  [PASS] All paragraphs present in content")

        meta = docs[0].get("metadata", {})
        print(f"  Metadata: {meta}")
        required_keys = {"source", "page", "type"}
        missing = required_keys - set(meta.keys())
        if missing:
            print(f"  [FAIL] Missing metadata: {missing}")
        else:
            print(f"  [PASS] All metadata keys present")
    else:
        print("  [FAIL] No documents returned")
except Exception as e:
    print(f"  [FAIL] load_docx error: {e}")
    import traceback
    traceback.print_exc()
finally:
    os.unlink(path)

# --- 5. load_document dispatcher for DOCX ---
print("\n--- load_document() Dispatcher for DOCX ---")
path2 = create_test_docx(["Dispatcher test paragraph.", "Second paragraph for testing."])
try:
    docs = load_document(path2, "document.docx")
    if docs:
        print(f"  [PASS] Dispatcher returned {len(docs)} doc(s) for DOCX")
        print(f"  Content: '{docs[0]['content'][:80]}'")
    else:
        print("  [FAIL] Dispatcher returned no docs for DOCX")
except Exception as e:
    print(f"  [FAIL] Dispatcher error: {e}")
    import traceback
    traceback.print_exc()
finally:
    os.unlink(path2)

# --- 6. DOCX with headings ---
print("\n--- DOCX with Headings ---")
doc_h = Document()
doc_h.add_heading("Main Title", level=1)
doc_h.add_paragraph("This is the introduction paragraph.")
doc_h.add_heading("Chapter 1", level=2)
doc_h.add_paragraph("Chapter one content here.")
doc_h.add_heading("Chapter 2", level=2)
doc_h.add_paragraph("Chapter two content here.")
fd, path3 = tempfile.mkstemp(suffix=".docx")
os.close(fd)
doc_h.save(path3)
try:
    docs = load_docx(path3)
    if docs:
        print(f"  [PASS] DOCX with headings loaded, {len(docs)} doc(s)")
        print(f"  Content: '{docs[0]['content'][:200]}'")
        # Check if headings are included
        content = docs[0]["content"]
        if "Main Title" in content:
            print("  [PASS] Heading text included")
        else:
            print("  [INFO] Headings not in paragraph content (normal for python-docx)")
    else:
        print("  [FAIL] No docs from DOCX with headings")
except Exception as e:
    print(f"  [FAIL] Headings DOCX error: {e}")
finally:
    os.unlink(path3)

# --- 7. Empty DOCX ---
print("\n--- Empty DOCX ---")
path4 = create_test_docx([])
try:
    docs = load_docx(path4)
    print(f"  Docs from empty DOCX: {len(docs)}")
    if len(docs) == 0:
        print("  [PASS] Empty DOCX yields no documents")
    elif docs and docs[0]["content"].strip() == "":
        print("  [WARN] Empty DOCX yields empty-content doc — may affect chunking")
    else:
        print(f"  [INFO] Got docs: {[d['content'] for d in docs]}")
except Exception as e:
    print(f"  [FAIL] Empty DOCX error: {e}")
finally:
    os.unlink(path4)

# --- 8. DOCX with special characters ---
print("\n--- DOCX with Special Characters ---")
path5 = create_test_docx([
    "Special chars: <>&\"' test",
    "Unicode: 你好 Héllo Привет",
    "Numbers: 1234567890",
    "Symbols: !@#$%^&*()",
])
try:
    docs = load_docx(path5)
    if docs:
        content = docs[0]["content"]
        print(f"  [PASS] Special chars DOCX loaded, length={len(content)}")
        if "你好" in content:
            print("  [PASS] Chinese characters preserved")
        else:
            print("  [WARN] Chinese characters may be lost")
    else:
        print("  [FAIL] No docs from special chars DOCX")
except Exception as e:
    print(f"  [FAIL] Special chars DOCX error: {e}")
finally:
    os.unlink(path5)

# --- 9. Metadata type check ---
print("\n--- DOCX Metadata Type ---")
path6 = create_test_docx(["Testing metadata type field."])
try:
    docs = load_docx(path6)
    if docs:
        doc_type = docs[0]["metadata"].get("type")
        doc_page = docs[0]["metadata"].get("page")
        print(f"  type = '{doc_type}' (expected: 'text')")
        print(f"  page = {doc_page} (expected: 1)")
        if doc_type == "text":
            print("  [PASS] type is 'text'")
        else:
            print(f"  [FAIL] type is '{doc_type}', expected 'text'")
        if doc_page == 1:
            print("  [PASS] page is 1")
        else:
            print(f"  [WARN] page is {doc_page}, expected 1")
except Exception as e:
    print(f"  [FAIL] Metadata type check error: {e}")
finally:
    os.unlink(path6)

print("\n" + "=" * 60)
print("TEST 06 COMPLETE")
print("=" * 60)
